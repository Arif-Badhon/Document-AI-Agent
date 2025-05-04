from docx import Document

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text.append(" | ".join(row_text))
    return "\n".join(text)

def chunk_text(text, chunk_size=500, overlap=100):
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)
        i += chunk_size - overlap
    return chunks